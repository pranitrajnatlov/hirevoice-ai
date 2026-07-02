"""
Structured resume parser (spec #1-6) — deterministic, no LLM in the hot path.

Pipeline:  extract_layout()  ->  detect_sections()  ->  field extraction  ->  schema

Tolerant of single/multi-page, multi-column, tables, varied fonts, and arbitrary
section ordering. Every extracted field carries a confidence (spec #4); low-confidence
fields are flagged in ``_meta`` rather than dropped. ``pdfplumber`` gives column-aware
PDF extraction with a graceful ``pypdf`` fallback; ``rapidfuzz`` tolerates heading noise.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from app.resume_integration import _extract_docx, _extract_pdf, load_resume_from_file
from app.vocabulary import fuzzy_match, normalize_skills

logger = logging.getLogger(__name__)

# ── Section heading aliases → canonical section key (spec #1) ──────────────────
SECTION_ALIASES: dict[str, list[str]] = {
    "experience": [
        "experience", "work experience", "professional experience", "employment history",
        "work history", "career history", "employment", "professional background",
    ],
    "skills": [
        "skills", "technical skills", "core competencies", "technologies", "tech stack",
        "skill set", "competencies", "areas of expertise", "key skills",
    ],
    "education": [
        "education", "academic qualification", "academic qualifications", "academics",
        "educational background", "academic background",
    ],
    "projects": [
        "projects", "personal projects", "academic projects", "key projects",
        "selected projects", "notable projects",
    ],
    "certifications": [
        "certifications", "certificates", "licenses & certifications", "licenses",
        "certification", "certifications & professional development",
        "professional development",
    ],
    "summary": [
        "summary", "professional summary", "profile", "about", "objective",
        "career objective", "about me", "professional profile",
    ],
    "languages": ["languages", "language proficiency", "languages known"],
    "personal_information": [
        "personal information", "contact", "contact information", "personal details",
        "contact details",
    ],
    "achievements": [
        "achievements", "awards", "honors", "awards & achievements", "accomplishments",
        "awards and achievements",
    ],
    "additional_information": [
        "additional information", "other information", "additional details",
        "other details", "miscellaneous",
    ],
    "internships": ["internships", "internship", "internship experience"],
    "volunteer": ["volunteer", "volunteer experience", "volunteering", "community service"],
    "publications": ["publications", "papers", "research", "research papers"],
    "interests": ["interests", "hobbies", "interests & hobbies", "interests and hobbies"],
}

_HEADING_CONF_EXACT = 0.95
_HEADING_FUZZY_CUTOFF = 80.0
_LOW_CONF_THRESHOLD = 0.6

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
_LINK_RE = re.compile(
    r"(https?://\S+|(?:www\.)?(?:linkedin|github)\.com/\S+|[a-z0-9-]+\.(?:linkedin|github)\.com/\S+|"
    r"(?:portfolio|blog|website)[.][a-z0-9.-]+\.[a-z]{2,}(?:/\S*)?)",
    re.I,
)
_YEAR_RANGE_RE = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec\w*)?\.?\s*"
    r"(?:19|20)\d{2})\s*[-–—to]+\s*(present|current|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec\w*)?\.?\s*(?:19|20)\d{2})",
    re.I,
)
_DEGREE_RE = re.compile(
    r"\b(bachelor|master|b\.?tech|m\.?tech|b\.?e\b|m\.?e\b|b\.?sc|m\.?sc|bca|mca|mba|ph\.?d|"
    r"diploma|intermediate|high school|associate)\b",
    re.I,
)
_ROLE_RE = re.compile(
    r"\b(developer|engineer|manager|analyst|consultant|designer|architect|intern|"
    r"lead|specialist|administrator|scientist|director|officer)\b",
    re.I,
)
_BULLET_RE = re.compile(r"^\s*[-•*▪◦‣·]\s+")
_SKILL_SPLIT_RE = re.compile(r"[,|;/]|\s{2,}|•|·|•|\(cid:\d+\)")

# Category prefixes commonly found in skills sections (e.g. "FRONTEND: React.js")
_SKILL_CATEGORY_PREFIX_RE = re.compile(
    r"^\s*(?:FRONTEND|BACKEND|DATABASES?|DEVOPS(?:\s*&\s*CLOUD)?|TOOLS(?:\s*&\s*PRACTICES)?|"
    r"CLOUD|LANGUAGES?|FRAMEWORKS?|LIBRARIES|OTHER|MOBILE|TESTING|ML|AI|DATA|INFRASTRUCTURE|"
    r"WEB|PROGRAMMING|SOFTWARE)\s*[:.]?\s*",
    re.I,
)
# Pattern that detects parenthesized sub-items: e.g. "AWS (EC2, S3, Lambda)"
_PAREN_SUB_RE = re.compile(r"([A-Za-z0-9./#+-]+(?:\s+[A-Za-z0-9./#+-]+)?)\s*\(([^)]+)\)")


# ── Layout extraction (spec #1, #3) ────────────────────────────────────────────
class ExtractResult:
    def __init__(self, lines: list[str], fmt: str, pages: int, raw_text: str):
        self.lines = lines
        self.format = fmt
        self.pages = pages
        self.raw_text = raw_text


def extract_layout(path: str | Path) -> ExtractResult:
    """Extract ordered text lines from PDF/DOCX/TXT, preserving reading order."""
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        lines, pages = _extract_pdf_layout(path)
        raw = "\n".join(lines)
        return ExtractResult(lines, "pdf", pages, raw)

    if suffix in (".docx", ".doc"):
        text = _extract_docx_layout(path)
        lines = _clean_lines(text)
        return ExtractResult(lines, "docx", 1, text)

    text = load_resume_from_file(path)  # TXT or unknown
    return ExtractResult(_clean_lines(text), "txt", 1, text)


def extract_from_text(text: str) -> ExtractResult:
    """Build an ExtractResult from already-extracted text (e.g. stored ``extracted_text``)."""
    return ExtractResult(_clean_lines(text), "text", 1, text)


def _extract_pdf_layout(path: Path) -> tuple[list[str], int]:
    """Column-aware PDF extraction via pdfplumber; falls back to pypdf plain text."""
    try:
        import pdfplumber  # type: ignore

        all_lines: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text(layout=True, keep_blank_chars=True)
                if text:
                    all_lines.extend(text.split("\n"))
        if all_lines:
            return _clean_lines("\n".join(all_lines)), pages
    except Exception as exc:  # pragma: no cover - depends on file/lib
        logger.warning("pdfplumber failed (%s) — falling back to pypdf", exc)

    text = _extract_pdf(path)  # existing pypdf path
    return _clean_lines(text), text.count("\f") + 1


def _extract_docx_layout(path: Path) -> str:
    """DOCX paragraphs + table cells (tables are common in resumes, spec #3)."""
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("  ".join(cells))
        return "\n".join(parts)
    except ImportError:  # pragma: no cover
        return _extract_docx(path)


def _clean_lines(text: str) -> list[str]:
    lines = []
    for ln in text.replace("\r", "\n").split("\n"):
        ln = ln.strip()
        # Remove common standalone footer "Page X" or "Page X of Y"
        if re.match(r"^Page\s+\d+(?:\s+of\s+\d+)?$", ln, re.I):
            continue
        if ln:
            lines.append(ln)
    return lines


# ── Section detection (spec #1) ────────────────────────────────────────────────
def _match_heading(line: str) -> Optional[tuple[str, float]]:
    """Return (canonical_section, confidence) if the line is a section heading."""
    raw = line.strip()
    words = raw.split()
    # Increase word/length limits to catch headings like "CERTIFICATIONS & PROFESSIONAL DEVELOPMENT"
    if not raw or len(words) > 7 or len(raw) > 60:
        return None
    if _EMAIL_RE.search(raw) or _PHONE_RE.search(raw):
        return None
    cleaned = re.sub(r"[^a-z& ]", "", raw.lower()).strip()
    if not cleaned:
        return None

    for canon, aliases in SECTION_ALIASES.items():
        if cleaned in aliases:
            return canon, _HEADING_CONF_EXACT
    # fuzzy — tolerate OCR / spacing noise; only for short heading-like lines
    flat_aliases = {a: c for c, al in SECTION_ALIASES.items() for a in al}
    hit = fuzzy_match(cleaned, flat_aliases.keys(), cutoff=_HEADING_FUZZY_CUTOFF)
    if hit:
        return flat_aliases[hit[0]], round(0.7 + (hit[1] - _HEADING_FUZZY_CUTOFF) / 100.0, 2)
    return None


def detect_sections(lines: list[str]) -> tuple[dict[str, list[str]], dict[str, float], list[str]]:
    """
    Bucket lines under detected canonical sections (order-independent, spec #1).

    Returns (sections, section_confidence, preamble) where preamble is everything
    before the first heading (usually name/contact/summary).
    """
    sections: dict[str, list[str]] = {}
    confidence: dict[str, float] = {}
    preamble: list[str] = []
    current: Optional[str] = None

    for line in lines:
        head = _match_heading(line)
        if head:
            current = head[0]
            sections.setdefault(current, [])
            confidence[current] = max(confidence.get(current, 0.0), head[1])
            continue
        if current is None:
            preamble.append(line)
        else:
            sections[current].append(line)

    return sections, confidence, preamble


# ── Field extraction ───────────────────────────────────────────────────────────
def _vc(value, confidence: float) -> dict:
    return {"value": value, "confidence": round(confidence, 2)}


def _is_headerish(line: str) -> bool:
    if _BULLET_RE.match(line):
        return False
    words = line.split()
    if not (1 <= len(words) <= 8):
        return False
    alpha = [w for w in words if any(c.isalpha() for c in w)]
    if not alpha:
        return False
    capped = sum(1 for w in alpha if w[0].isupper() or w.isupper())
    return capped / len(alpha) >= 0.6 and not line.rstrip().endswith(".")


def _extract_contact(preamble: list[str], full_text: str, base_conf: float) -> dict:
    info: dict = {}
    email = _EMAIL_RE.search(full_text)
    links = _LINK_RE.findall(full_text)
    if email:
        info["email"] = _vc(email.group(0), 0.95)
    # Phone: require >= 10 digits so date ranges like "2020 - 07" don't match.
    for m in _PHONE_RE.finditer(full_text):
        if len(re.sub(r"\D", "", m.group(0))) >= 10:
            info["phone"] = _vc(m.group(0).strip(), 0.8)
            break
    if links:
        info["links"] = [l if isinstance(l, str) else l[0] for l in links][:5]

    # Name: first short, title-case, alpha-only line near the top that isn't a
    # tech/skill line (reject commas and known technologies like "Spring Boot, Java").
    name, name_conf = None, 0.0
    for i, line in enumerate(preamble[:15]):
        words = line.split()
        if not (2 <= len(words) <= 4) or "," in line:
            continue
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line) or any(ch.isdigit() for ch in line):
            continue
        if _match_heading(line):
            continue
        if re.search(r"\b(ltd|limited|inc|pvt|llc|technologies|solutions|corp|company|university|college|institute)\b", line, re.I):
            continue
        from app.vocabulary import normalize_skill
        if any(normalize_skill(w)[1] >= 0.9 for w in words):
            continue
        if all(w[0].isupper() or w.isupper() for w in words if w):
            name = line.strip()
            name_conf = 0.7 if i <= 3 else 0.6
            break
    if name:
        info["name"] = _vc(name, name_conf)
    return info


def _extract_experience(lines: list[str], base_conf: float) -> list[dict]:
    entries: list[dict] = []
    cur: Optional[dict] = None

    def _flush():
        nonlocal cur
        if cur and (cur.get("company") or cur.get("title") or cur.get("responsibilities")):
            entries.append(cur)
        cur = None

    def _is_cid_bullet(line: str) -> bool:
        """Detect (cid:NNN) bullet markers from pdfplumber."""
        return bool(re.match(r"^\s*\(cid:\d+\)", line))

    for line in lines:
        # Clean up (cid:NNN) bullet markers
        clean_line = re.sub(r"\(cid:\d+\)", "", line).strip()
        if not clean_line:
            continue

        date_hit = _YEAR_RANGE_RE.search(clean_line)
        bullet = bool(_BULLET_RE.match(clean_line)) or _is_cid_bullet(line)
        headerish = _is_headerish(clean_line) and not bullet

        # Handle pipe-delimited lines: "Company Name | Location" or "Company | Location | Date"
        pipe_parts = [p.strip() for p in clean_line.split("|") if p.strip()]
        is_company_location_line = (
            len(pipe_parts) >= 2
            and not bullet
            and not _ROLE_RE.search(clean_line)
            and any(c.isalpha() for c in pipe_parts[0])
        )

        if headerish and cur and cur.get("responsibilities"):
            _flush()  # new entry begins after we've collected responsibilities
        if cur is None:
            cur = {"company": None, "title": None, "dates": None, "location": None,
                   "responsibilities": [], "technologies": [], "confidence": round(base_conf, 2)}

        if is_company_location_line and cur["company"] is None:
            # Parse "Company | Location" or "Company | Location | Date" lines
            cur["company"] = pipe_parts[0]
            if len(pipe_parts) >= 2:
                # Check if any part contains a date
                for part in pipe_parts[1:]:
                    date_in_part = _YEAR_RANGE_RE.search(part)
                    if date_in_part:
                        cur["dates"] = date_in_part.group(0).strip()
                    else:
                        cur["location"] = part
        elif date_hit and not bullet:
            cur["dates"] = date_hit.group(0).strip()
        elif _ROLE_RE.search(clean_line) and cur["title"] is None and headerish:
            cur["title"] = clean_line.strip()
        elif headerish and cur["company"] is None and not is_company_location_line:
            cur["company"] = clean_line.strip()
        elif bullet or not headerish:
            text = _BULLET_RE.sub("", clean_line).strip()
            if text:
                cur["responsibilities"].append(text)

    _flush()

    # derive technologies per entry from its responsibility text
    for e in entries:
        joined = " ".join(e["responsibilities"])
        techs = _detect_inline_tech(joined)
        e["technologies"] = techs
    return entries


def _extract_education(lines: list[str], base_conf: float) -> list[dict]:
    entries: list[dict] = []
    for line in lines:
        if _DEGREE_RE.search(line) or re.search(r"\b(university|college|institute|school)\b", line, re.I):
            degree = _DEGREE_RE.search(line)
            entries.append({
                "institution": line.strip() if re.search(r"\b(university|college|institute|school)\b", line, re.I) else None,
                "degree": degree.group(0) if degree else None,
                "dates": (_YEAR_RANGE_RE.search(line).group(0) if _YEAR_RANGE_RE.search(line) else None),
                "raw": line.strip(),
                "confidence": round(base_conf, 2),
            })
    return entries


def _extract_projects(lines: list[str], base_conf: float) -> list[dict]:
    entries: list[dict] = []
    cur: Optional[dict] = None

    def _is_project_header(line: str) -> bool:
        """True if line looks like a project title (short, capitalized, non-bullet).
        Rejects comma-separated tech lists which look headerish but aren't titles."""
        if _BULLET_RE.match(line):
            return False
        clean = re.sub(r"\(cid:\d+\)", "", line).strip()
        if not clean:
            return False
        # Reject if it looks like a lone technology (e.g. "Vercel", "AWS", "TanStack Query")
        from app.vocabulary import normalize_skill
        if len(clean.split()) <= 2:
            _, conf = normalize_skill(clean)
            if conf >= 0.85:
                return False
        # Reject comma-heavy lines (tech lists)
        if clean.count(",") >= 2:
            return False
        return _is_headerish(clean) and len(clean.split()) <= 12

    for line in lines:
        clean = re.sub(r"\(cid:\d+\)", "", line).strip()
        if not clean:
            continue

        # Handle "Project Name | Tech1, Tech2, ..." on one line
        if "|" in clean:
            parts = clean.split("|", 1)
            name_part = parts[0].strip().rstrip("|")
            tech_part = parts[1].strip() if len(parts) > 1 else ""

            # Only treat as a new project header if the name part looks like a title
            if name_part and _is_headerish(name_part) and len(name_part.split()) >= 3:
                if cur:
                    entries.append(cur)
                techs = _detect_inline_tech(tech_part) if tech_part else []
                cur = {"name": name_part, "description": "", "technologies": techs,
                       "confidence": round(base_conf, 2)}
                continue

        if _is_project_header(clean) and len(clean.split()) >= 3:
            if cur:
                entries.append(cur)
            cur = {"name": clean.strip(), "description": "", "technologies": [],
                   "confidence": round(base_conf, 2)}
        elif cur is not None:
            text = _BULLET_RE.sub("", clean).strip()
            cur["description"] = (cur["description"] + " " + text).strip()
    if cur:
        entries.append(cur)
    for e in entries:
        # Merge inline tech from description with those already found from the header
        desc_techs = _detect_inline_tech(e["description"])
        existing = set(e["technologies"])
        e["technologies"] = e["technologies"] + [t for t in desc_techs if t not in existing]
    return entries


def _extract_list_items(lines: list[str]) -> list[str]:
    items: list[str] = []
    for line in lines:
        text = _BULLET_RE.sub("", line).strip()
        for part in _SKILL_SPLIT_RE.split(text):
            part = part.strip()
            if len(part) >= 2:
                items.append(part)
    return items


_SKILL_STOPWORDS = {
    "and", "the", "i", "a", "an", "with", "on", "for", "to", "of", "in", "we",
    "furthermore", "also", "my", "our", "using", "used", "key", "description",
}


def _looks_like_skill(token: str) -> bool:
    """Reject prose fragments that leak into a skills bucket (1-4 short words, no sentence markers)."""
    words = token.split()
    if not (1 <= len(words) <= 4):
        return False
    if token.endswith((".", ":")) or ". " in token:
        return False
    if words[0].lower() in _SKILL_STOPWORDS:
        return False
    return any(any(c.isalpha() for c in w) for w in words)


def _extract_skills(lines: list[str]) -> list[dict]:
    """Skill candidates filtered to plausible skills, then normalized to canonical forms."""
    # Pre-process: strip category prefixes (FRONTEND:, BACKEND:, etc.) and expand
    # parenthesized sub-items (e.g. "AWS (EC2, S3, Lambda)" -> "AWS", "EC2", "S3", "Lambda")
    cleaned_lines: list[str] = []
    for line in lines:
        # Remove (cid:NNN) markers
        line = re.sub(r"\(cid:\d+\)", " ", line)
        # Strip category prefixes
        line = _SKILL_CATEGORY_PREFIX_RE.sub("", line)
        # Expand parenthesized sub-skills: "AWS (EC2, S3, Lambda, RDS)" -> "AWS, EC2, S3, Lambda, RDS"
        expanded = line
        for m in _PAREN_SUB_RE.finditer(line):
            parent = m.group(1).strip()
            children = [c.strip() for c in m.group(2).split(",") if c.strip()]
            replacement = ", ".join([parent] + children)
            expanded = expanded.replace(m.group(0), replacement)
        cleaned_lines.append(expanded)

    candidates = [c for c in _extract_list_items(cleaned_lines) if _looks_like_skill(c)]
    skills = normalize_skills(candidates)
    # keep known tech (high conf) or genuinely short unknown tokens (<=2 words) as niche skills
    return [s for s in skills if s["confidence"] >= 0.85 or len(s["value"].split()) <= 2]


def _detect_inline_tech(text: str) -> list[str]:
    """Find known technologies mentioned inside free text (for experience/project tech lists)."""
    if not text:
        return []
    found = normalize_skills(_extract_list_items([text]))
    return [s["value"] for s in found if s["confidence"] >= 0.85]


# ── Additional Information parser ──────────────────────────────────────────────
def _parse_additional_information(lines: list[str]) -> dict:
    """
    Parse the 'Additional Information' composite section which typically contains
    mixed sub-fields like Languages, Interests, Achievements on one or more lines
    separated by pipes.
    """
    result: dict[str, list[str]] = {"languages": [], "achievements": [], "interests": []}
    full_text = " ".join(re.sub(r"\(cid:\d+\)", " ", ln) for ln in lines)

    # Clean up some common PDF artifacts before processing
    full_text = full_text.replace(" | ", "|").replace(" |", "|").replace("| ", "|")
    
    # Simple explicit string matching to find the start of each section
    lang_idx = max(full_text.lower().find("languages:"), full_text.lower().find("language:"))
    int_idx = max(full_text.lower().find("interests:"), full_text.lower().find("interest:"))
    ach_idx = max(full_text.lower().find("achievements:"), full_text.lower().find("awards:"))

    # Extract languages if present
    if lang_idx != -1:
        # Find where the next section starts, or end of string, or the pipe character if it's acting as a major section delimiter
        end_idx = len(full_text)
        for idx in [int_idx, ach_idx]:
            if idx != -1 and idx > lang_idx and idx < end_idx:
                end_idx = idx
        
        # Also look for a pipe before the next section
        pipe_idx = full_text.find("|", lang_idx)
        if pipe_idx != -1 and pipe_idx < end_idx:
            # If there's a pipe, it often delimits major sections like "Languages: English | Interests: ..."
            # But sometimes it delimits items. If the text after pipe doesn't contain another key soon, it might be an item.
            # Let's assume pipe acts as an item delimiter here just like comma.
            pass

        lang_str = full_text[lang_idx:end_idx].split(":", 1)[1].strip()
        # Clean trailing pipe if it's right before the next section
        if lang_str.endswith("|"):
            lang_str = lang_str[:-1]
            
        items = re.split(r"[|•·,]|\(cid:\d+\)", lang_str)
        result["languages"].extend([i.strip() for i in items if len(i.strip()) >= 2])

    # Extract interests if present
    if int_idx != -1:
        end_idx = len(full_text)
        for idx in [lang_idx, ach_idx]:
            if idx != -1 and idx > int_idx and idx < end_idx:
                end_idx = idx
                
        int_str = full_text[int_idx:end_idx].split(":", 1)[1].strip()
        if int_str.endswith("|"):
            int_str = int_str[:-1]
            
        items = re.split(r"[|•·,]|\(cid:\d+\)", int_str)
        result["interests"].extend([i.strip() for i in items if len(i.strip()) >= 2])

    # Extract achievements if present
    if ach_idx != -1:
        end_idx = len(full_text)
        for idx in [lang_idx, int_idx]:
            if idx != -1 and idx > ach_idx and idx < end_idx:
                end_idx = idx
                
        ach_str = full_text[ach_idx:end_idx].split(":", 1)[1].strip()
        if ach_str.endswith("|"):
            ach_str = ach_str[:-1]
            
        # For achievements, we don't want to split by comma as much because it breaks up sentences
        items = re.split(r"[|•·]|\(cid:\d+\)", ach_str)
        result["achievements"].extend([i.strip() for i in items if len(i.strip()) >= 2])

    return result


# ── Orchestration (spec #2, #5) ────────────────────────────────────────────────
def parse_resume(source: str | Path, *, is_text: bool = False) -> dict:
    """
    Parse a resume file (or already-extracted text) into the structured schema.

    Returns the schema dict from spec #2 plus a ``_meta`` block with parser confidence,
    detected sections, and low-confidence field flags.
    """
    extract = extract_from_text(str(source)) if is_text else extract_layout(source)
    sections, sec_conf, preamble = detect_sections(extract.lines)

    def conf(name: str, default: float = 0.5) -> float:
        return sec_conf.get(name, default)

    # Skills (most important downstream — drives vocabulary + interview context)
    skills = _extract_skills(sections.get("skills", []))

    # Certifications: clean (cid:NNN) bullet markers from items
    cert_lines = sections.get("certifications", [])
    cert_items = _extract_list_items([re.sub(r"\(cid:\d+\)", " ", ln) for ln in cert_lines])

    # Additional information (composite section with languages, achievements, interests)
    addl_info = _parse_additional_information(sections.get("additional_information", []))

    # Merge languages/achievements from dedicated sections + additional_information
    lang_items = _extract_list_items(sections.get("languages", [])) + addl_info["languages"]
    achv_items = _extract_list_items(sections.get("achievements", [])) + addl_info["achievements"]

    profile: dict = {
        "personal_information": _extract_contact(preamble, extract.raw_text, conf("personal_information", 0.7)),
        "summary": _vc(" ".join(sections.get("summary", []))[:600], conf("summary", 0.5)) if sections.get("summary") or preamble else _vc("", 0.0),
        "experience": _extract_experience(sections.get("experience", []) + sections.get("internships", []), conf("experience", 0.6)),
        "education": _extract_education(sections.get("education", []), conf("education", 0.6)),
        "skills": skills,
        "projects": _extract_projects(sections.get("projects", []), conf("projects", 0.6)),
        "certifications": cert_items,
        "languages": lang_items,
        "achievements": achv_items,
        "interests": addl_info["interests"],
    }

    # If no explicit summary heading, use the longest preamble sentence as a soft summary.
    if not profile["summary"]["value"] and preamble:
        longest = max(preamble, key=len)
        if len(longest) > 60:
            profile["summary"] = _vc(longest[:600], 0.5)

    # ── Confidence aggregation + low-confidence flags (spec #4) ──
    low_conf: list[str] = []
    for s in skills:
        if s["confidence"] < _LOW_CONF_THRESHOLD:
            low_conf.append(f"skills.{s['value']}")
    pi = profile["personal_information"]
    for k, v in pi.items():
        if isinstance(v, dict) and v.get("confidence", 1) < _LOW_CONF_THRESHOLD:
            low_conf.append(f"personal_information.{k}")

    section_confs = list(sec_conf.values()) or [0.4]
    parser_conf = round(sum(section_confs) / len(section_confs), 2)

    profile["_meta"] = {
        "parser_confidence": parser_conf,
        "sections_found": sorted(sections.keys()),
        "low_confidence_fields": low_conf,
        "format": extract.format,
        "pages": extract.pages,
        "skill_count": len(skills),
    }
    return profile
